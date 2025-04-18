import pandas as pd
pd.set_option('future.no_silent_downcasting', True)

import os
from abc import ABC, abstractmethod
from langchain.docstore.document import Document as LangDocument

class sf_default:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        from langchain_community.document_loaders import PyPDFLoader
        from langchain_community.document_loaders import Docx2txtLoader
        from langchain.text_splitter import (
            RecursiveCharacterTextSplitter,
        )
        basename, extension = os.path.splitext(self.file_path)

        match extension:
            case ".docx":
                loader = Docx2txtLoader(self.file_path)
            case ".pdf":  
                loader = PyPDFLoader(self.file_path)
            case _:
                print(f"Данный файл не поддерживается {self.file_path}")
                return []

        documents = loader.load()

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=200,)
        documents = text_splitter.split_documents(documents)     

        return documents
    
class sf_DataProcessing:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        from langchain.text_splitter import (
            RecursiveCharacterTextSplitter,
        )

        loader = sfDocumentLoaderFactory.create_loader(self.file_path) 
        documents = loader.load_documents() 
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50) 
        documents = text_splitter.split_documents(documents)
        for i, doc in enumerate(documents[:3]):
            print(f"Чанк {i}: {doc.page_content[:500]}")

        return documents

class sf_add_keywords_512_chunk:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        import os
        import re 
        from langchain_community.document_loaders import PyPDFLoader
        from langchain_community.document_loaders import Docx2txtLoader
        from langchain.text_splitter import (
            RecursiveCharacterTextSplitter,
        )
    
        # читаем документ
        basename, extension = os.path.splitext(self.file_path)
        match extension:
            case ".docx":
                loader = Docx2txtLoader(self.file_path)
            case ".pdf":  
                loader = PyPDFLoader(self.file_path)
            case _:
                print(f"Данный файл не поддерживается {self.file_path}")
                return []
        documents = loader.load()

        # Объявляем класс
        doc_c = get_keywords(documents)
        # находим слова
        # ВАЖНО! слова находятся через сеть, необходимо  установить сеть deepseek-r1:latest
        # или заменить llm_class и llm_keywords
        keywords = doc_c.get_keywords_def()
        # в keywords у нас хранятся ключевые слова
        print (keywords)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100,)
        chunks = text_splitter.split_documents(documents)

        # в chunk должна быть итоговый класс, мы просто добавляем в каждый чанк ключевые слова
        enriched_chunks = [doc_c.enrich_chunk_with_additional_info(chunk, keywords) for chunk in chunks]

        # и возращаем этот чанк
        return enriched_chunks

class sf_DataProcessing_keywords_512_chunk:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        import os
        import re 
        from langchain_community.document_loaders import PyPDFLoader
        from langchain_community.document_loaders import Docx2txtLoader
        from langchain.text_splitter import (
            RecursiveCharacterTextSplitter,
        )
    
        # читаем документ
        from langchain.text_splitter import (
            RecursiveCharacterTextSplitter,
        )

        loader = sfDocumentLoaderFactory.create_loader(self.file_path) 
        documents = loader.load_documents() 

        # Объявляем класс
        doc_c = get_keywords(documents)
        # находим слова
        # ВАЖНО! слова находятся через сеть, необходимо  установить сеть deepseek-r1:latest
        # или заменить llm_class и llm_keywords
        keywords = doc_c.get_keywords_def()
        # в keywords у нас хранятся ключевые слова
        print (keywords)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100,)
        chunks = text_splitter.split_documents(documents)

        # в chunk должна быть итоговый класс, мы просто добавляем в каждый чанк ключевые слова
        enriched_chunks = [doc_c.enrich_chunk_with_additional_info(chunk, keywords) for chunk in chunks]

        # и возращаем этот чанк
        return enriched_chunks
    
class sf_DataProcessing_keywords_512_chunk_and_Tables:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        from langchain.text_splitter import (
            RecursiveCharacterTextSplitter,
        )
    
        # читаем документ
        from langchain.text_splitter import (
            RecursiveCharacterTextSplitter,
        )

        loader = sfDocumentLoaderFactory.create_loader(self.file_path) 
        documents = loader.load_documents() 

        # # Объявляем класс
        doc_c = get_keywords(documents)
        # # находим слова
        # # ВАЖНО! слова находятся через сеть, необходимо  установить сеть deepseek-r1:latest
        # # или заменить llm_class и llm_keywords
        keywords = doc_c.get_keywords_def()
        # # в keywords у нас хранятся ключевые слова
        print (keywords)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50,)
        chunks = text_splitter.split_documents(documents)

        # в chunk должна быть итоговый класс, мы просто добавляем в каждый чанк ключевые слова
        enriched_chunks = [doc_c.enrich_chunk_with_additional_info(chunk, keywords) for chunk in chunks]

        # и возращаем этот чанк
        return enriched_chunks

# Абстрактный базовый класс для загрузчиков документов
class sfBaseDocumentLoader(ABC):
    @abstractmethod
    def load_documents(self) -> list:
        # Метод должен вернуть список объектов LangDocument
        pass

# Класс для определения расширения файла
class sfFileTypeDetector:
    @staticmethod
    def get_file_type(file_path: str) -> str:
        _, ext = os.path.splitext(file_path)
        return ext.lower()

# # Класс для извлечения изображений из DOCX и применения OCR        
# class sfDOCXImageExtractor:
#     def __init__(self, file_path: str):
#         self.file_path = file_path

#     def extract_images_text(self) -> str:
#         import zipfile
#         import pytesseract
#         from PIL import Image

#         ocr_text = []
#         try:
#             with zipfile.ZipFile(self.file_path) as docx_zip:
#                 for file in docx_zip.namelist():
#                     if file.startswith("word/media/"):
#                         try:
#                             image_data = docx_zip.read(file)
#                             image = Image.open(io.BytesIO(image_data))
#                             if image.format == 'WMF':
#                                 image = image.convert('RGB') # Конвертируем WMF в поддерживаемый pytesseract формат 
#                             text = pytesseract.image_to_string(image).strip()
#                             if text:
#                                 ocr_text.append(text)
#                         except Exception as img_e:
#                             print(f"Ошибка обработки изображения {file}: {img_e}")
#         except Exception as e:
#             print (f"Ошибка открытия DOCX как zip-архива: {e}")
#         return "\n\n".join(ocr_text)
            
# Фабрика для обработки DOCX (извлечение текста, таблиц, изображения
class sfDOCXLoader(sfBaseDocumentLoader):
    def __init__(self, file_path: str, loader_type: str):
        self.file_path = file_path
        self.loader_type = loader_type

    def clean_text(self, text:str) -> str:
        import re
        
        # Удаление невидимых символов
        invisible_chars = ['\u200b', '\ufeff', '\xa0', '\x0c']
        for char in invisible_chars:
            text = text.replace(char, ' ' if char == '\xa0' else '')

        # Замена повторяющихся символов: -----, ===, **** и т.п.
        text = re.sub(r'([\-=_*~#]{3,})', '', text)

        # Удаление лишних пробелов с сохранением абзацев
        # 1. сначала нормализуем пробелы внутри строк
        text = re.sub(r'[ \t]+', ' ', text)
        # 2. заменяем 3 и более переносов строк на 2 (абзац)
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 3. удаляем пробелы в начале и конце строк
        text = "\n".join([line.strip() for line in text.splitlines()])
        return text.strip()

    def load_documents(self):
        if self.loader_type == "unstructured":
            from  langchain_community.document_loaders import UnstructuredWordDocumentLoader
            loader = UnstructuredWordDocumentLoader(self.file_path)
        elif self.loader_type == "Docx2txtLoader":
            from langchain_community.document_loaders import Docx2txtLoader
            loader = Docx2txtLoader(self.file_path)
        else:
            raise ValueError(f"Неизвестный тип загрузчика для DOCX: {self.loader_type}")
        
        documents = loader.load()

        # Очищаем документ
        for doc in documents:
            doc.page_content = self.clean_text(doc.page_content)

        # Обновляем метаданные, чтобы source содержал только имя файла
        file_name = os.path.basename(self.file_path)
        for doc in documents:
            doc.metadata["source"] = file_name

        # Дополнительно извлекаем таблицы
        tables_text, table_metadata = self.extract_tables()
        if tables_text.strip():
            documents.append(LangDocument(page_content=tables_text, metadata=table_metadata))

        return documents
    
    def extract_tables(self):
        from docx import Document as DocxDocument
        extracted_tables = []
        doc = DocxDocument(self.file_path)
        try:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append("\t".join(cells))
                table_text = "\n".join(rows)
                if table_text.strip():
                    extracted_tables.append(table_text)
        except Exception as e:
            print(f"Ошибка извлечения таблиц: {e}")

        if extracted_tables:
            tables_text = "\n\n".join(extracted_tables)
        else:
            tables_text = ""
        table_metadata = {"source": os.path.basename(self.file_path)}
        return tables_text, table_metadata

# Класс для загрузки PDF, объединяющий текст и таблицы
class sfPDFLoader(sfBaseDocumentLoader):
    def __init__(self, file_path: str, loader_type: str):
        self.file_path = file_path
        self.loader_type = loader_type
    
    def load_documents(self):
        # Выбор загрузчика для файла
        if self.loader_type == "py":
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(self.file_path)
        elif self.loader_type == "unstructured":
            from langchain_community.document_loaders import UnstructuredPDFLoader
            loader = UnstructuredPDFLoader(self.file_path)
        elif self.loader_type == "plumber":
            from langchain_community.document_loaders import PDFMinerLoader
            loader = PDFMinerLoader(self.file_path)        
        else:
            raise ValueError(f"Неизвестный тип загрузчика для PDF: {self.loader_type}")
        
        documents = loader.load() 

        # Обновляем метаданные, чтобы source содержал только имя файла
        file_name = os.path.basename(self.file_path)
        for doc in documents:
            doc.metadata["source"] = file_name

        # Дополнительно извлекаем таблицы
        tables_text, table_metadata = self.extract_tables()
        if tables_text.strip():
            documents.append(LangDocument(page_content=tables_text, metadata=table_metadata))

        return documents

    def extract_tables(self):
        extracted_tables = []
        metadata = {"source": os.path.basename(self.file_path)}

        import fitz  # PyMuPDF
        try:
            doc = fitz.open(self.file_path)
        except Exception as e:
            print(f"Ошибка открытия PDF: {e}")
            return ""
        
        for page in doc:
            page_num = page.number + 1
            analysis = self.analyze_page(page)
            if analysis["contains_tables"]:
                table_text = self.try_extract_table(page, page_num)
                if table_text:
                    extracted_tables.append(table_text)
        return "\n\n".join(extracted_tables) if extracted_tables else "", metadata        


    def analyze_page(self, page):
        # Анализируем есть ли на странице таблица
        analysis = {"contains_tables": False}
        try:
            tables = page.find_tables()
            if tables.tables: 
                analysis["contains_tables"] = True
        except Exception as e:
            print(f"Анализ страницы {page} завершился с ошибкой: {e}")
        return analysis

    def try_extract_table(self, page, page_num):
        import camelot
        table_text_candidates = {}

        try:
            page_dict = page.get_text("dict")
            contains_lines = any("lines" in block for block in page_dict.get("blocks", []))
            flavors = ["stream", "hybrid"] if contains_lines else ["stream"]

            for flavor in flavors:
                try:
                    tables = camelot.read_pdf(self.file_path, pages=str(page_num), flavor=flavor)
                    if tables.n > 0:
                        quality = sum(len(table.df) for table in tables)
                        table_text = "\n".join([table.df.to_csv(index=False) for table in tables])
                        table_text_candidates[flavor] = (quality, table_text)
                except Exception as e:
                    print(f"Flavor '{flavor}' не сработал на странице {page_num}: {e}")

            if not table_text_candidates:
                print(f"Таблицы не извлечены на странице {page_num}.")
                self._cache[page_num] = ""
                return ""
                        
            best_flavor = max(table_text_candidates, key=lambda f: table_text_candidates[f][0])
            best_quality, best_text = table_text_candidates[best_flavor]
            print(f"Выбран режим '{best_flavor}' для страницы {page_num} с качеством {best_quality}.")
            return best_text
        except Exception as e:
            print(f"Ошибка извлечения таблиц: {e}")
            return ""

class get_keywords:
    
    from typing import List
    from langchain_ollama import OllamaLLM
    
    promt_category = {
        "Закон или нормативный акт": (
            {"Номер закона или нормативного акта " : "Какой номер документа?",
             "Наименование закона или нормативного акта " : "Какое наименование документа?",
             "Дата закона или нормативного акта " : "Какая дата документа?"}
        ),
        "Договор": (
            {"Номер договора ":"Какой номер договора?",
             "Дата заключения договора ":"Какая дата договора?",
             "Договор заключен между ":" Между кем заключен договор?",
             "Предмет договора ":"Какой предмет договора?"}
        ),
        "Письмо": (
            {"Письмо от ":"Кто написал письмо?",
            "Дата письма ":"Какая дата письма?",
            "Получатель письма " : "Кто получатель письма?",
            "Тема письма " : "Какая тема письма ?"}
        ),
        "Курсовая или дипломная работа": (
            {"Тема курсовой или дипломной работы ":"Какая тема работы?",
            "Автор курсовой или дипломной работы ":"Кто подготовил работу?"}
        ),
        "Информация об организации": (
            {"Наименование организации ":" Какое наименование организации?"}
        ),
        "Техническое задание или технический проект": (
            {"Наименование технического задания ":"Какое наименование документа?",
            "Номер технического задания ":"Какой номер документа?",
            "Автор технического задания ": "Кто подготовил документ?"}
        ),
        "Рассказ или повесть": (
            {"Автор художественного произведения ": "Кто автор документа?", 
            "Название произведения": "Какое название документа?"}
        ),
        "Неизвестная категория": (
            {}
        )
    }
    X_char = 1000
    llm_class = OllamaLLM(model="deepseek-r1:latest", temperature = "0.1")
    llm_keywords = OllamaLLM(model="llama3:latest", temperature = "0.0")
    
    def __init__(self, documents: List[LangDocument]):
        self.documents = documents
    
    def remove_text_between_tags(self, text: str):
        start_tag = '<think>'
        end_tag = '</think>'

        result = []
        last_position = 0

        while True:
            start_index = text.find(start_tag, last_position)
            if start_index == -1:
                break

            result.append(text[last_position:start_index])

            end_index = text.find(end_tag, start_index + len(start_tag))
            if end_index == -1:
                break

            last_position = end_index + len(end_tag)

        result.append(text[last_position:])

        return ''.join(result)
    
        
    def clean_doc(self):
        """
        Заменяет все символы перевода строки в page_content у каждого документа.
        """
        for doc in self.documents:
            doc.page_content = doc.page_content.replace('\n', '')

    def get_X_characters(self) -> str:
        """
        Возвращает X символов из списка документов, объединяя page_content.
        """
#         result = ""
#         for doc in self.documents:
#             if len(result) >= self.X_char:
#                 break
#             result += doc.page_content[:self.X_char - len(result)]
#         return result
        result_start = ""
        result_end = ""

        # Собираем X символов с начала списка
        for doc in self.documents:
            if len(result_start) >= self.X_char:
                break
            result_start += doc.page_content[:self.X_char - len(result_start)]

        # Собираем X символов с конца списка (начиная с конца)
        for doc in reversed(self.documents):
            if len(result_end) >= self.X_char:
                break
            result_end = doc.page_content[-(self.X_char - len(result_end)):] + result_end

        return result_start + result_end

    def define_document_type(self, doc_for_context: str) -> str:
        """
        Определение типа документа.
        """
        category_str = ""

        for category, promt in self.promt_category.items():
            category_str = category_str + category + ", "
        promt_category = f"""Контекст: мы проводим работы по классификации текстов, необходимо определять к какой категории относится текст
        Роль: твоя роль по части текста определять к какой из предложенных категории относится этот текст 
        Задача: Тебе предоставлен текст {doc_for_context} ты должен определить к какой  
        категорий из списка он относится, список категорий: {category_str}. 
        Критерии Качества: необходимо предоставить точно одну из предоставленных списка категорий, нельзя менять использовать другие слова, должно быть только название категории"""
        llm_response = self.llm_class.invoke(promt_category)
        return self.remove_text_between_tags(llm_response)

            
    def find_category(self, text: str) -> str:
        """
        Определяет категорию текста на основе словаря promt_category.
        """
        for category in self.promt_category.keys():
            if category.lower() in text.lower():
                return category
        return "Неизвестная категория"
    
    def return_promt_find_keywords(self, doc_type: str) -> str:
        """
        Возвращает промт для поиска ключевых слов.
        """
        promt_ = ""
        dictionary = self.promt_category
        if doc_type in dictionary:
            promt_ =  dictionary[doc_type]
        
        return promt_
    
    def find_keywords(self, input_dic, doc_char: str) -> str:
        """
        Возвращает ключевые слова.
        """ 
        promt_keyword = "Вы полезный ассистент. Вы отвечаете на вопросы о документации, используя эти данные: {self.data}. Ответь на русском языке на этот запрос: {self.prompt} "
        promt_keyword = promt_keyword.replace("{self.data}", doc_char)
            
        #parts = input_str.split(";")  # Разделяем строку по ";"
        modified_parts = []  # Создаем список для измененных частей

        for key in input_dic:
            promt_llm = promt_keyword.replace("{self.prompt}", input_dic[key])
            llm_response = self.llm_keywords.invoke(promt_llm)
            modified_parts.append(key + " " + llm_response)  

        return ";".join(modified_parts)  # Объединяем обратно в строку


        #llm_response = self.llm_pd.invoke(promt_keyword)
        #return self.remove_text_between_tags(llm_response)
    
    def add_keywords(self, additional_info: str):
        """
        Добавляет ключевое слово в page_content каждого документа под ключом 'keywords'.
        """
        for doc in self.documents:
            if 'page_content' not in doc.__dict__:
                doc.page_content = ""
            doc.page_content += f"\n\n{additional_info}"
    
    def get_keywords_def(self):
        self.clean_doc()
        doc_char = self.get_X_characters()
        doc_type_llm = self.define_document_type(doc_char)
        doc_type = self.find_category(doc_type_llm)
        promt_key = self.return_promt_find_keywords(doc_type)
        return self.find_keywords(promt_key, doc_char)
    
    def enrich_chunk_with_additional_info(self, doc, additional_text):
        """Добавляет additional_info в текст чанка, но не объединяет все метаданные"""
        enriched_content = f"{additional_text}\n\n{doc.page_content}"  # 👈 Добавляем только `additional_info`
        return LangDocument(page_content=enriched_content, metadata=doc.metadata)

# Фабрика загрузчиков, которая определяет тип файла и возвращает нужный загрузчик
class sfDocumentLoaderFactory:
    @staticmethod
    def create_loader(file_path: str):
        ext = sfFileTypeDetector.get_file_type(file_path)
        if ext == ".pdf":
            return sfPDFLoader(file_path, loader_type="py")
        elif ext == ".docx":
            return sfDOCXLoader(file_path, loader_type="Docx2txtLoader")
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")

#  Конвейер обработки документа: загрузка -> разделение на чанки
class sfDocumentProcessingPipeline:
    def __init__(self, file_path: str):
        self.loader = sfDocumentLoaderFactory.create_loader(file_path)
        # self.splitter = sfTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    
    def separate_file(self):
        documents = self.loader.load_documents()
        return self.splitter.split(documents)