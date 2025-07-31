# имена классов должны начинаться с sf_
# и отображать основные характеристики (по усмотрению разработчика)
import pandas as pd
pd.set_option('future.no_silent_downcasting', True)
import os, re, inspect
import fitz
from abc import ABC, abstractmethod
from langchain.docstore.document import Document as LangDocument
from tabulate import tabulate  # ✅ Импорт добавлен
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangDocument
import re
from langchain.text_splitter import CharacterTextSplitter

class sf_default:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        from langchain_community.document_loaders import PyPDFLoader
        from langchain_community.document_loaders import Docx2txtLoader
        # from langchain.text_splitter import (
        #     RecursiveCharacterTextSplitter,
        # )
        basename, extension = os.path.splitext(self.file_path)
        match extension:
            case ".docx":
                loader = Docx2txtLoader(self.file_path)
            case ".pdf":  
                loader = PyPDFLoader(self.file_path)
            case _:
                print(f"Данный файл не поддерживается {self.file_path}")
                return []
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=200,)
        documents = text_splitter.split_documents(documents)     
        return documents

class sf_DataProcessing:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        # from langchain.text_splitter import (
        #     RecursiveCharacterTextSplitter,
        # )
        loader = sfDocumentLoaderFactory.create_loader(self.file_path) 
        documents = loader.load_documents() 
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50) 
        documents = text_splitter.split_documents(documents)
        for i, doc in enumerate(documents[:3]):
            print(f"Чанк {i}: {doc.page_content[:500]}")
        return documents

class sf_add_keywords_512_chunk:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        from langchain_community.document_loaders import PyPDFLoader
        from langchain_community.document_loaders import Docx2txtLoader
        from langchain_community.document_loaders import UnstructuredPowerPointLoader
        # from langchain.text_splitter import (
        #     RecursiveCharacterTextSplitter,
        # )
        # читаем документ
        basename, extension = os.path.splitext(self.file_path)
        match extension:
            case ".docx":
                print("Statr load docx")
                loader = Docx2txtLoader(self.file_path)
            case ".pdf": 
                print("Statr load pdf") 
                loader = PyPDFLoader(self.file_path)
            case ".pptx":
                print("Statr load pptx")   
                loader = UnstructuredPowerPointLoader(self.file_path)
            case _:
                print(f"Данный файл не поддерживается {self.file_path}")
                return []
        documents = loader.load()
        # Объявляем класс
        doc_c = get_keywords(documents)
        # находим слова
        # ВАЖНО! слова находятся через сеть, необходимо  установить сеть deepseek-r1:latest
        # или заменить llm_class и llm_keywords
        keywords = doc_c.get_keywords_def()
        # в keywords у нас хранятся ключевые слова
        print (keywords)
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100,)
        chunks = text_splitter.split_documents(documents)
        # в chunk должна быть итоговый класс, мы просто добавляем в каждый чанк ключевые слова
        enriched_chunks = [doc_c.enrich_chunk_with_additional_info(chunk, keywords) for chunk in chunks]
        # и возращаем этот чанк
        return enriched_chunks

class sf_DataProcessing_keywords_512_chunk:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        from langchain_community.document_loaders import PyPDFLoader
        from langchain_community.document_loaders import Docx2txtLoader
        # from langchain.text_splitter import (
        #     RecursiveCharacterTextSplitter,
        # )
        # читаем документ
        # from langchain.text_splitter import (
        #     RecursiveCharacterTextSplitter,
        # )
        loader = sfDocumentLoaderFactory.create_loader(self.file_path) 
        documents = loader.load_documents() 
        # Объявляем класс
        doc_c = get_keywords(documents)
        # находим слова
        # ВАЖНО! слова находятся через сеть, необходимо  установить сеть deepseek-r1:latest
        # или заменить llm_class и llm_keywords
        keywords = doc_c.get_keywords_def()
        # в keywords у нас хранятся ключевые слова
        print (keywords)
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100,)
        chunks = text_splitter.split_documents(documents)
        # в chunk должна быть итоговый класс, мы просто добавляем в каждый чанк ключевые слова
        enriched_chunks = [doc_c.enrich_chunk_with_additional_info(chunk, keywords) for chunk in chunks]
        # и возращаем этот чанк
        return enriched_chunks

class sf_DataProcessing_keywords_512_chunk_and_Tables:
    def __init__(self, file_path):
        self.file_path = file_path
    def separate_file(self):
        from langchain.text_splitter import (
            RecursiveCharacterTextSplitter,
        )
        # читаем документ
        # from langchain.text_splitter import (
        #     RecursiveCharacterTextSplitter,
        # )
        loader = sfDocumentLoaderFactory.create_loader(self.file_path) 
        documents = loader.load_documents() 
        print("Данные documents")
        print(len(documents))
        if len(documents) == 0:
            documents =  [LangDocument(
                    page_content="ERROR: файл содержит только изображения или не содержит текста",
                    metadata={"source": self.file_path, "type": "error", "error": "no_text_or_tables"})]
        # # Объявляем класс
        doc_c = get_keywords(documents)
        # # находим слова
        # # ВАЖНО! слова находятся через сеть, необходимо  установить сеть deepseek-r1:latest
        # # или заменить llm_class и llm_keywords
        keywords = doc_c.get_keywords_def()
        # # в keywords у нас хранятся ключевые слова
        print (keywords)

        # text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100,)
        # chunks = text_splitter.split_documents(documents)
        # chunks = fix_broken_tables(chunks)

        # Создаём умный сплиттер
        splitter = SmartTextSplitter(default_chunk_size=512)

        # Разбиваем документы
        chunks = splitter.split_documents(documents)

        # в chunk должна быть итоговый класс, мы просто добавляем в каждый чанк ключевые слова
        #enriched_chunks = [doc_c.enrich_chunk_with_additional_info(chunk, keywords) for chunk in chunks]
        # ...existing code...
        # в chunk должна быть итоговый класс, мы просто добавляем в каждый чанк ключевые слова
        enriched_chunks = []
        for i, chunk in enumerate(chunks):
            #print(f"[LOG] Чанк {i}: {chunk.page_content[:1500]}")  # Логируем первые 500 символов чанка
            enriched_chunk = doc_c.enrich_chunk_with_additional_info(chunk, keywords)
            enriched_chunks.append(enriched_chunk)
        # ...existing code...
        # и возращаем этот чанк
        return enriched_chunks

################################
################################ Вспомогательные классы
################################

def fix_broken_tables(chunks):
    i = 0
    fixed = []
    while i < len(chunks):
        doc = chunks[i]

        # Если началась таблица, но не закончилась
        if "[TABLE_START]" in doc.page_content and "[TABLE_END]" not in doc.page_content:
            new_content = doc.page_content
            j = i + 1
            while j < len(chunks):
                new_content += chunks[j].page_content
                if "[TABLE_END]" in chunks[j].page_content:
                    break
                j += 1
            # Соединяем
            fixed.append(LangDocument(
                page_content=new_content,
                metadata=doc.metadata
            ))
            i = j + 1
        else:
            fixed.append(doc)
            i += 1
    return fixed

class SmartTextSplitter:
    def __init__(self, default_chunk_size=512, overlap=100):
        self.default_chunk_size = default_chunk_size
        self.overlap = overlap

    def split_documents(self, documents):
        result = []
        for doc in documents:
            content = doc.page_content
            metadata = doc.metadata

            # Если это таблица — добавляем как есть
            if metadata.get("type") == "table":
                result.append(doc)
            else:
                # Объединяем текст без разрыва по строкам
                clean_text = " ".join(content.split())
                chunks = self._split_text_by_chunk_size(clean_text)
                for i, chunk in enumerate(chunks):
                    new_metadata = metadata.copy()
                    new_metadata["chunk"] = i
                    result.append(LangDocument(page_content=chunk, metadata=new_metadata))
        return result

    def _split_text_by_chunk_size(self, text: str):
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.default_chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start += self.default_chunk_size - self.overlap
        return chunks
    
# class SmartTextSplitter:
#     def __init__(self, default_chunk_size=512, overlap=100, table_tag_start="[TABLE_START]", table_tag_end="[TABLE_END]"):
#         self.default_chunk_size = default_chunk_size
#         self.overlap = overlap
#         self.table_tag_start = table_tag_start
#         self.table_tag_end = table_tag_end
#         self.default_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=default_chunk_size,
#             chunk_overlap=overlap,
#             separators=["\n\n", "\n", " ", ""]
#         )

#     def split_documents(self, documents):
#         """
#         Принимает список LangDocument (текст и таблицы),
#         Возвращает список LangDocument, где:
#           - Таблицы всегда целиком в одном чанке
#           - Обычный текст делится на части
#         """
#         result = []
#         for doc in documents:
#             content = doc.page_content
#             metadata = doc.metadata

#             # Если это таблица — добавляем как есть
#             if metadata.get("type") == "table" or \
#                (self.table_tag_start in content and self.table_tag_end in content):
#                 result.append(doc)
#             else:
#                 # Иначе разбиваем на чанки
#                 chunks = self.default_splitter.split_text(content)
#                 for i, chunk in enumerate(chunks):
#                     new_metadata = metadata.copy()
#                     new_metadata["chunk"] = i
#                     result.append(LangDocument(page_content=chunk, metadata=new_metadata))
#         return result

# Фабрика загрузчиков, которая определяет тип файла и возвращает нужный загрузчик
class sfDocumentLoaderFactory:
    @staticmethod
    def create_loader(file_path: str):
        ext = sfFileTypeDetector.get_file_type(file_path)
        if ext == ".pdf":
            return sfPDFLoader(file_path)
        elif ext == ".docx":
            return sfDOCXLoader(file_path, loader_type="python-docx")
        elif ext == ".pptx":  
            return sfPPTXLoader(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")

#  Конвейер обработки документа: загрузка -> разделение на чанки
class sfDocumentProcessingPipeline:
    def __init__(self, file_path: str):
        self.loader = sfDocumentLoaderFactory.create_loader(file_path)
        # self.splitter = sfTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    def separate_file(self):
        documents = self.loader.load_documents()
        return self.splitter.split(documents)

class sfBaseDocumentLoader(ABC):
    @abstractmethod
    def load_documents(self) -> list:
        # Метод должен вернуть список объектов LangDocument
        pass

# Класс для определения расширения файла
class sfFileTypeDetector:
    @staticmethod
    def get_file_type(file_path: str) -> str:
        _, ext = os.path.splitext(file_path)
        return ext.lower()

class sfDOCXLoader(sfBaseDocumentLoader):
    def __init__(self, file_path: str, loader_type: str):
        self.file_path = file_path
        self.loader_type = loader_type

    def clean_text(self, text:str) -> str:
        import re
        invisible_chars = ['\u200b', '\ufeff', '\xa0', '\x0c']
        for char in invisible_chars:
            text = text.replace(char, ' ' if char == '\xa0' else '')
        text = re.sub(r'([\-=_*~#]{3,})', '', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n', text)
        text = "\n".join([line.strip() for line in text.splitlines()])
        return text.strip()

    def extract_blocks(self):
        from docx import Document
        from docx.table import Table
        doc = Document(self.file_path)
        blocks = []
        for child in doc.element.body:
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                texts = [n.text for n in child.iter() if n.tag.endswith('}t') and n.text]
                text = self.clean_text(''.join(texts))
                if text:
                    blocks.append(('paragraph', text))
            elif tag == 'tbl':
                tbl = Table(child, doc)
                headers = []
                rows = []
                for i, row in enumerate(tbl.rows):
                    cells = [self.clean_text(cell.text) for cell in row.cells]
                    if i == 0:
                        headers = cells  # Первая строка — заголовки
                    else:
                        rows.append(cells)
                if not headers or not rows:
                    continue

                # Преобразуем в Markdown-таблицу
                markdown_table = tabulate(rows, headers=headers, tablefmt='github')

                # Добавляем разделители, чтобы таблица не разрывалась на чанки
                full_table = "[TABLE_START]\n" + markdown_table + "\n[TABLE_END]"

                blocks.append(('table', full_table))
        return blocks

    def extract_tables(self):
        try:
            from docx2python import docx2python
            result = docx2python(self.file_path)
            table_docs = []
            table_idx = 1
            for section in result.body:
                for element in section:
                    if isinstance(element, list) and element and all(isinstance(row, list) for row in element):
                        table_rows = []
                        for row in element:
                            clean_cells = [cell.strip().replace('\n', ' ') for cell in row]
                            if any(clean_cells):
                                table_rows.append(clean_cells)
                        if not table_rows:
                            continue
                        header = table_rows[0]
                        data_rows = table_rows[1:]
                        lines = ["\t".join(header)]
                        for row in data_rows:
                            lines.append("\t".join(row))
                        table_text = f"Таблица {table_idx}:\n" + "\n".join(lines)
                        table_docs.append(table_text)
                        table_idx += 1
            tables_text = "\n".join(table_docs)
            table_metadata = {"source": os.path.basename(self.file_path)}
            print("\n=== Итоговый результат для RAG ===\n")
            print(tables_text)
            return tables_text, table_metadata
        except Exception as e:
            print(f"Ошибка при обработке таблиц ({self.loader_type}): {e}")
            import traceback
            traceback.print_exc()
            return []

    def load_documents(self):
        try:
            docs = []
            blocks = self.extract_blocks()

            # Список для хранения обычного текста
            current_text = []

            for kind, text in blocks:
                if kind == "table":
                    # Если был накопленный текст — добавляем его как отдельный документ
                    if current_text:
                        full_text = "\n".join(current_text)
                        docs.append(LangDocument(
                            page_content=full_text,
                            metadata={"source": os.path.basename(self.file_path), "type": "paragraph"}
                        ))
                        current_text = []

                    # Добавляем таблицу как отдельный документ
                    docs.append(LangDocument(
                        page_content=text,
                        metadata={"source": os.path.basename(self.file_path), "type": "table"}
                    ))
                else:
                    # Накапливаем обычный текст
                    current_text.append(text)

            # Не забываем про оставшийся текст после последней таблицы
            if current_text:
                full_text = "\n".join(current_text)
                docs.append(LangDocument(
                    page_content=full_text,
                    metadata={"source": os.path.basename(self.file_path), "type": "paragraph"}
                ))

        except Exception as e:
            print(f"Ошибка при загрузке документа: {e}")
            return [LangDocument(
                page_content="ERROR: файл содержит только изображения или не содержит текста",
                metadata={"source": self.file_path, "type": "error", "error": "no_text_or_tables"})]

        return docs

class sfPDFLoader(sfBaseDocumentLoader):
    def __init__(
            self, 
            file_path: str, 
            flavor: str = "stream",
            hf_k: int=5, 
            hf_thr_ratio: float = 0.9 
            ):
        self.file_path = file_path
        self.flavor = flavor 
        self.k = hf_k
        self.thr = hf_thr_ratio

    def _clean(self, text: str) -> str:
        invisible = ["\u200b", "\ufeff", "\xa0", "\x0c"]
        for ch in invisible:
            text = text.replace(ch, " " if ch == "\xa0" else "")
        text = re.sub(r"[\t ]+", " ", text)
        text = re.sub(r"\n{3,}", "\n", text)
        return text.strip()

    def _bbox_to_rect(self, bbox: tuple[float, float, float, float], pad: float = 5) -> "fitz.Rect":
        x1, y1, x2, y2 = bbox 
        rect = fitz.Rect(x1, y1, x2, y2)
        if hasattr(rect, "inflate"):
            return rect.inflate(pad, pad)
        return fitz.Rect(x1 - pad, y1 - pad, x2 + pad, y2 + pad)

    @staticmethod
    def _norm(s: str) -> str:
        s = s.lower()
        s = re.sub(r"\d{1,4}[./-]\d{1,2}[./-]?\d{0,4}", "", s)
        s = re.sub(r"\d+", "", s)
        s = re.sub(r"\s+", " ", s)
        return s.strip()

    def _collect_hf_candidates(self, pdf: fitz.Document) -> tuple[set[str], set[str]]:
        k = self.k
        thr_ratio = self.thr
        head_cnt: dict[str, int] = {} 
        foot_cnt: dict[str, int] = {}
        for page in pdf:
            blocks = sorted(page.get_text("blocks"), key=lambda b: b[1])
            top_lines = [blocks[i][4].strip() for i in range(min(k, len(blocks)))]
            bot_lines = [blocks[-(i + 1)][4].strip() for i in range(min(k, len(blocks)))]
            for ln in top_lines:
                key = self._norm(ln)
                if key:
                    head_cnt[key] = head_cnt.get(key, 0) + 1
            for ln in bot_lines:
                key = self._norm(ln)
                if key:
                    foot_cnt[key] = foot_cnt.get(key, 0) + 1
        thresh = max(1, int(len(pdf) * thr_ratio)) 
        headers_set = {s for s, c in head_cnt.items() if c >= thresh}
        footers_set = {s for s, c in foot_cnt.items() if c >= thresh}
        return headers_set, footers_set

    def load_documents(self) -> list[LangDocument]:
        import camelot
        docs: list[LangDocument] = []
        pdf = fitz.open(self.file_path)
        file_name = os.path.basename(self.file_path)
        hdr_set, ftr_set = self._collect_hf_candidates(pdf)
        table_idx = 1
        for page in pdf:
            page_id = page.number + 1
            try:
                tables = camelot.read_pdf(
                    self.file_path,
                    pages=str(page_id),
                    flavor=self.flavor,
                    strip_text="\n",
                    split_text=True,
                    edge_tol=200,
                    row_tol=5,
                )
            except Exception as e:
                print(f"[Camelot] ошибка при разборке {page_id}: {e}")
                tables = []
            tbl_rects: list[fitz.Rect] = []
            for t in tables:
                tbl_rects.append(self._bbox_to_rect(t._bbox))
            try:
                blocks = page.get_text("blocks")
                text_parts: list[str] = []
                for x0, y0, x1, y1, txt, *_ in blocks:
                    rect = fitz.Rect(x0, y0, x1, y1)
                    if not txt.strip():
                        continue
                    if page_id != 1 and any(
                        self._norm(line) in hdr_set or self._norm(line) in ftr_set
                        for line in txt.splitlines()
                        if line.strip()
                    ):
                        continue
                    text_parts.append(txt)
            except Exception as e:
                print(e)
            if text_parts:
                para_text = self._clean("\n".join(text_parts))
                is_guess = ("\t" in para_text and para_text.count("\t") >= 3) \
                            or ("..." in para_text and para_text.count("...") >= 5)
                docs.append(
                    LangDocument(
                        page_content=para_text,
                        metadata={
                            "source": file_name,
                            "type": "table_guess" if is_guess else "paragraph",
                            "page": page_id,
                        },
                    )
                )
            tables = tables or []
            table_order = sorted(zip(tbl_rects, tables), key=lambda p: p[0].y0)
            for rect, table in table_order:
                if table.df.shape[0] < 3 or table.df.shape[1] < 2:
                    continue
                first_cell_norm = self._norm(table.df.iloc[0, 0])
                # ✅ Изменение: таблица в формате Markdown внутри разделителя
                markdown_table = tabulate(table.df.values.tolist(), headers=table.df.columns.tolist(), tablefmt="github")
                content = f"[TABLE_START]\nТаблица {table_idx}:\n{markdown_table}\n[TABLE_END]"
                docs.append(
                    LangDocument(
                        page_content=content,
                        metadata={
                            "source": file_name,
                            "type": "table",
                            "page": page_id,
                            "rows": table.df.shape[0],
                            "cols": table.df.shape[1],
                        },
                    )
                )
                table_idx += 1
        pdf.close()
        return docs

class sfPPTXLoader(sfBaseDocumentLoader):
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load_documents(self):
        from langchain_community.document_loaders import UnstructuredPowerPointLoader
        loader = UnstructuredPowerPointLoader(self.file_path)
        documents = loader.load()  # <-- ВАЖНО: вызываем .load()
        return documents

class get_keywords:
    import base64
    from app.config import settings_llm, settings
    from typing import List
    from langchain_ollama import OllamaLLM
    promt_category = {
        "Закон или нормативный акт": (
            {"Номер закона или нормативного акта " : "Какой номер документа?",
             "Наименование закона или нормативного акта " : "Какое наименование документа?",
             "Дата закона или нормативного акта " : "Какая дата документа?"}
        ),
        "Договор": (
            {"Номер договора ":"Какой номер договора?",
             "Дата заключения договора ":"Какая дата договора?",
             "Договор заключен между ":" Между кем заключен договор?",
             "Предмет договора ":"Какой предмет договора?"}
        ),
        "Письмо": (
            {"Письмо от ":"Кто написал письмо?",
            "Дата письма ":"Какая дата письма?",
            "Получатель письма " : "Кто получатель письма?",
            "Тема письма " : "Какая тема письма ?"}
        ),
        "Курсовая или дипломная работа": (
            {"Тема курсовой или дипломной работы ":"Какая тема работы?",
            "Автор курсовой или дипломной работы ":"Кто подготовил работу?"}
        ),
        "Информация об организации": (
            {"Наименование организации ":" Какое наименование организации?"}
        ),
        "Техническое задание или технический проект": (
            {"Наименование технического задания ":"Какое наименование документа?",
            "Номер технического задания ":"Какой номер документа?",
            "Автор технического задания ": "Кто подготовил документ?"}
        ),
        "Рассказ или повесть": (
            {"Автор художественного произведения ": "Кто автор документа?", 
            "Название произведения": "Какое название документа?"}
        ),
        "Неизвестная категория": (
            {}
        )
    }
    encoded_credentials = base64.b64encode(f"{settings_llm.USER_LLM}:{settings_llm.PASSWORD_LLM}".encode()).decode()
    headers = {'Authorization': f'Basic {encoded_credentials}'}
    X_char = 1000
    llm_class = OllamaLLM( 
                model="gemma3:12b", temperature = 0.1, base_url=settings_llm.URL_LLM, client_kwargs={'headers': headers})
    llm_keywords = OllamaLLM( 
                model="gemma3:12b", temperature = 0.0, base_url=settings_llm.URL_LLM, client_kwargs={'headers': headers})
    def __init__(self, documents: List[LangDocument]):
        self.documents = documents
    def remove_text_between_tags(self, text: str):
        start_tag = '</think>'
        end_tag = '</think>'
        result = []
        last_position = 0
        while True:
            start_index = text.find(start_tag, last_position)
            if start_index == -1:
                break
            result.append(text[last_position:start_index])
            end_index = text.find(end_tag, start_index + len(start_tag))
            if end_index == -1:
                break
            last_position = end_index + len(end_tag)
        result.append(text[last_position:])
        return ''.join(result)
    def clean_doc(self):
        for doc in self.documents:
            doc.page_content = doc.page_content.replace('\n', '')
    def get_X_characters(self) -> str:
        result_start = ""
        result_end = ""
        for doc in self.documents:
            if len(result_start) >= self.X_char:
                break
            result_start += doc.page_content[:self.X_char - len(result_start)]
        for doc in reversed(self.documents):
            if len(result_end) >= self.X_char:
                break
            result_end = doc.page_content[-(self.X_char - len(result_end)):] + result_end
        return result_start + result_end
    def define_document_type(self, doc_for_context: str) -> str:
        category_str = ""
        for category, promt in self.promt_category.items():
            category_str = category_str + category + ", "
        promt_category = f"""Контекст: мы проводим работы по классификации текстов, необходимо определять к какой категории относится текст
        Роль: твоя роль по части текста определять к какой из предложенных категории относится этот текст 
        Задача: Тебе предоставлен текст {doc_for_context} ты должен определить к какой  
        категорий из списка он относится, список категорий: {category_str}. 
        Критерии Качества: необходимо предоставить точно одну из предоставленных списка категорий, нельзя менять использовать другие слова, должно быть только название категории"""
        llm_response = self.llm_class.invoke(promt_category)
        return self.remove_text_between_tags(llm_response)
    def find_category(self, text: str) -> str:
        for category in self.promt_category.keys():
            if category.lower() in text.lower():
                return category
        return "Неизвестная категория"
    def return_promt_find_keywords(self, doc_type: str) -> str:
        promt_ = ""
        dictionary = self.promt_category
        if doc_type in dictionary:
            promt_ =  dictionary[doc_type]
        return promt_
    def find_keywords(self, input_dic, doc_char: str) -> str:
        promt_keyword = "Вы полезный ассистент. Вы отвечаете на вопросы о документации, используя эти данные: {self.data}. Ответь на русском языке на этот запрос: {self.prompt} "
        promt_keyword = promt_keyword.replace("{self.data}", doc_char)
        modified_parts = []  
        for key in input_dic:
            promt_llm = promt_keyword.replace("{self.prompt}", input_dic[key])
            llm_response = self.llm_keywords.invoke(promt_llm)
            modified_parts.append(key + " " + llm_response)  
        return ";".join(modified_parts)  
    def add_keywords(self, additional_info: str):
        for doc in self.documents:
            if 'page_content' not in doc.__dict__:
                doc.page_content = ""
            doc.page_content += f"\n{additional_info}"
    def get_keywords_def(self):
        self.clean_doc()
        doc_char = self.get_X_characters()
        doc_type_llm = self.define_document_type(doc_char)
        doc_type = self.find_category(doc_type_llm)
        promt_key = self.return_promt_find_keywords(doc_type)
        return self.find_keywords(promt_key, doc_char)
    def enrich_chunk_with_additional_info(self, doc, additional_text):
        enriched_content = f"{additional_text}\n{doc.page_content}"  
        return LangDocument(page_content=enriched_content, metadata=doc.metadata)
